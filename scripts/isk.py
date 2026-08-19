#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Генератор искового заявления о взыскании в порядке регресса 1/2 доли
платежей по ипотечному кредиту + процентов по ст. 395 ГК РФ.

Проценты считаются автоматически по ключевой ставке Банка России
(официальный SOAP-сервис DailyInfo, scripts/cbr_rate.py).
Госпошлина — по шкале ст. 333.19 НК РФ (с 08.09.2024).

Запуск:
    .venv/bin/python scripts/isk.py

Значения из .env (в корне проекта) подставляются как умолчания:
Enter — принять, ручной ввод — переопределить. Формат ключей — .env.example.
"""

import re
from datetime import date, timedelta

from cbr_rate import CbrUnavailableError, days_in_year, keyrate_series
from pretenziya import (
    ask, ask_date, ask_int, ask_money, ask_payments, confirm, env_default,
    load_env, money, parse_env_payments, valid_date,
)


def to_date(s: str) -> date:
    d, m, y = map(int, s.split("."))
    return date(y, m, d)


# ---------------------------------------------------------------------------
# Госпошлина: ст. 333.19 НК РФ, шкала с 08.09.2024 (ФЗ № 259-ФЗ)
# ---------------------------------------------------------------------------

def court_fee(price: float) -> float:
    if price <= 100_000:
        return 4_000.0
    if price <= 200_000:
        return 4_000 + 0.03 * (price - 100_000)
    if price <= 1_000_000:
        return 7_000 + 0.025 * (price - 200_000)
    if price <= 3_000_000:
        return 27_000 + 0.01 * (price - 1_000_000)
    if price <= 8_000_000:
        return 47_000 + 0.007 * (price - 3_000_000)
    if price <= 20_000_000:
        return 107_000 + 0.005 * (price - 8_000_000)
    return min(207_000 + 0.003 * (price - 20_000_000), 400_000)


# ---------------------------------------------------------------------------
# Проценты по ст. 395 ГК РФ
# ---------------------------------------------------------------------------

def calc_interest(payments: list, doc_date_str: str,
                  start_mode: str, claim_delivered: str) -> tuple:
    """Возвращает ([(дата, платёж, 1/2 доли, проценты)], итого процентов).

    start_mode: '1' — со дня, следующего за датой каждого платежа;
                '2' — со дня, следующего за датой вручения претензии
                      (но не ранее дня, следующего за датой платежа).
    Ставка — ключевая ставка ЦБ на каждый день просрочки (день/год).
    """
    doc_date = to_date(doc_date_str)
    items = []
    for pdate_str, amount in payments:
        pdate = to_date(pdate_str)
        start = pdate + timedelta(days=1)
        if start_mode == "2" and claim_delivered:
            start = max(start, to_date(claim_delivered) + timedelta(days=1))
        end = doc_date
        if end < start:
            items.append((pdate_str, amount, amount / 2, 0.0))
            continue
        series = keyrate_series(start, end)
        interest = 0.0
        for day, rate in series:
            interest += (amount / 2) * (rate / 100) / days_in_year(day.year)
        items.append((pdate_str, amount, amount / 2, round(interest, 2)))
    return items, round(sum(i[3] for i in items), 2)


def ask_interest(payments: list, doc_date: str, claim_delivered: str) -> tuple:
    """Расчёт процентов с выбором даты начала + офлайн-fallback."""
    while True:
        mode = ask(
            "Начало начисления процентов по ст. 395: 1 — со дня после каждого "
            "платежа (рекомендуется) / 2 — со дня после вручения претензии", "1").strip()
        if mode in ("1", "2"):
            break
        print("  Введите 1 или 2.")
    try:
        return calc_interest(payments, doc_date, mode, claim_delivered)
    except CbrUnavailableError as e:
        print(f"\n  ⚠ Не удалось получить ключевую ставку ЦБ: {e}")
        print("  Варианты: ввести единую ставку вручную (грубая оценка ко всему периоду)")
        print("  или 0 — сформировать иск без расчётной суммы процентов.")
        rate = None
        while rate is None:
            raw = input("Единая ставка, % годовых (или 0): ").strip().replace(",", ".")
            try:
                rate = float(raw)
                if rate < 0:
                    raise ValueError
            except ValueError:
                print("  Введите неотрицательное число, например 16 или 16.5")
                rate = None
        if rate == 0:
            return None, 0
        items = []
        for pdate_str, amount in payments:
            pdate = to_date(pdate_str)
            start = pdate + timedelta(days=1)
            if mode == "2" and claim_delivered:
                start = max(start, to_date(claim_delivered) + timedelta(days=1))
            days = (to_date(doc_date) - start).days + 1
            # грубая оценка единой ставкой: 365 дней без учёта високосности
            intr = round((amount / 2) * (rate / 100) / 365 * max(days, 0), 2)
            items.append((pdate_str, amount, amount / 2, intr))
        return items, round(sum(i[3] for i in items), 2)


# ---------------------------------------------------------------------------
# Сбор данных
# ---------------------------------------------------------------------------

def collect_common(env: dict) -> dict:
    """Общие поля профиля (совместимы с pretenziya.py)."""
    d = {}
    print("--- Отправитель/истец (вы) ---")
    d["sender_name"] = ask("Ваше ФИО полностью", env_default(env, "SENDER_NAME"))
    d["sender_address"] = ask("Ваш адрес регистрации (с индексом)",
                              env_default(env, "SENDER_ADDRESS"))
    d["sender_phone"] = ask("Ваш телефон", env_default(env, "SENDER_PHONE"))

    print("\n--- Ответчик (бывшая супруга) ---")
    d["recipient_name"] = ask("ФИО полностью", env_default(env, "RECIPIENT_NAME"))
    d["recipient_address"] = ask("Адрес регистрации (с индексом)",
                                 env_default(env, "RECIPIENT_ADDRESS"))

    print("\n--- Брак ---")
    d["marriage_date"] = ask_date("Дата заключения брака",
                                  env_default(env, "MARRIAGE_DATE", valid_date))
    d["marriage_cert"] = ask("Свидетельство о заключении брака (серия, номер)",
                             env_default(env, "MARRIAGE_CERT"))
    d["divorce_date"] = ask_date("Дата расторжения брака",
                                 env_default(env, "DIVORCE_DATE", valid_date))
    d["divorce_cert"] = ask("Свидетельство о расторжении брака (серия, номер) "
                            "или суд + дата решения",
                            env_default(env, "DIVORCE_CERT"))

    print("\n--- Кредитный договор ---")
    d["bank_name"] = ask("Наименование банка", env_default(env, "BANK_NAME"))
    d["contract_num"] = ask("Номер кредитного договора", env_default(env, "CONTRACT_NUM"))
    d["contract_date"] = ask_date("Дата кредитного договора",
                                  env_default(env, "CONTRACT_DATE", valid_date))
    d["sender_account"] = ask("Ваш счёт, с которого платите",
                              env_default(env, "SENDER_ACCOUNT"))
    d["apartment_address"] = ask("Адрес заложенной квартиры",
                                 env_default(env, "APARTMENT_ADDRESS"))
    d["unpaid_from"] = ask_date("Дата, с которой она перестала возмещать долю",
                                env_default(env, "UNPAID_FROM", valid_date))

    print("\n--- Платежи ---")
    preset, bad = parse_env_payments(env.get("PAYMENTS", ""))
    if bad:
        print(f"  ⚠ В PAYMENTS из .env пропущены некорректные записи: {'; '.join(bad)}")
    d["payments"] = ask_payments(preset, not_before=d["divorce_date"])
    return d


def collect_claim(env: dict) -> dict:
    """Поля искового заявления (суд, идентификаторы, досудебный порядок)."""
    c = {}
    print("\n--- Суд ---")
    c["court"] = ask("Наименование суда в дательном/предложном падеже "
                     "(«Мировому судье судебного участка №…» / «В N-ский "
                     "районный суд»)",
                     env_default(env, "COURT_NAME"))
    price_hint = ("цена иска до 50 000 ₽ — мировой судья (п. 4 ч. 1 ст. 23 ГПК); "
                  "выше — районный суд (порог 100 000 ₽ действует только для "
                  "потребительских споров)")
    print(f"  Подсказка: {price_hint}.")

    print("\n--- Истец: обязательные реквизиты (ч. 2 ст. 131 ГПК) ---")
    c["plaintiff_birth"] = ask_date("Ваша дата рождения",
                                    env_default(env, "PLAINTIFF_BIRTH_DATE", valid_date))
    c["plaintiff_id"] = ask("Ваш паспорт (серия, номер, кем и когда выдан) — "
                            "идентификатор истца",
                            env_default(env, "PLAINTIFF_PASSPORT"))

    print("\n--- Ответчик: если известно (иначе суд запросит) ---")
    c["defendant_birth"] = ask("Дата рождения ответчика (или «-», если неизвестна)",
                               env_default(env, "DEFENDANT_BIRTH_DATE", valid_date, "-"))
    c["defendant_id"] = ask("Паспорт/ИНН/СНИЛС ответчика (или «-»)",
                            env_default(env, "DEFENDANT_IDENTIFIER", None, "-"))

    print("\n--- Досудебный порядок (претензия) ---")
    c["claim_sent"] = ask_date("Дата отправки претензии заказным письмом",
                               env_default(env, "CLAIM_SENT_DATE", valid_date))
    c["claim_track"] = ask("Трек-номер (РПО)", env_default(env, "CLAIM_TRACK"))
    c["claim_delivered"] = ask("Дата вручения претензии (или «не вручена»)",
                               env_default(env, "CLAIM_DELIVERED_DATE"))
    if not valid_date(c["claim_delivered"]):
        c["claim_delivered"] = ""
    c["claim_result"] = ""
    while c["claim_result"] not in ("1", "2"):
        c["claim_result"] = ask("Результат: 1 — без ответа / 2 — отказ платить",
                                env_default(env, "CLAIM_RESULT", None, "1"))
        if c["claim_result"] not in ("1", "2"):
            print("  Введите 1 или 2.")
    return c


# ---------------------------------------------------------------------------
# Документ
# ---------------------------------------------------------------------------

def build_doc(d: dict, c: dict, interest_items, interest_total: float,
              doc_date: str, path: str) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    debt = sum(a for _, a in d["payments"]) / 2
    price = debt + (interest_total or 0)
    fee = court_fee(price)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.5
    for section in doc.sections:
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    def para(text, align=None, bold=False, space_after=6):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        if align:
            p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        return p

    # Шапка
    para(c["court"], space_after=12)
    para("Истец:", bold=True, space_after=0)
    para(f"{d['sender_name']}, {c['plaintiff_birth']} г.р., "
         f"паспорт {c['plaintiff_id']},", space_after=0)
    para(f"адрес: {d['sender_address']}, тел.: {d['sender_phone']}", space_after=12)
    para("Ответчик:", bold=True, space_after=0)
    db = (f", {c['defendant_birth']} г.р."
          if c["defendant_birth"] not in ("", "-") else "")
    di = (f", идентификатор: {c['defendant_id']}"
          if c["defendant_id"] not in ("", "-") else "")
    para(f"{d['recipient_name']}{db}{di},", space_after=0)
    para(f"адрес: {d['recipient_address']}", space_after=12)
    para(f"Цена иска: {money(price)} ₽", space_after=0)
    para(f"Государственная пошлина: {money(fee)} ₽", space_after=18)

    # Заголовок
    para("ИСКОВОЕ ЗАЯВЛЕНИЕ", WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=0)
    para("о взыскании в порядке регресса 1/2 доли платежей по кредитному",
         WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    para("договору и процентов за пользование чужими денежными средствами",
         WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

    delivered = (f"и вручена {c['claim_delivered']}" if c["claim_delivered"]
                 else "по данным Почты России не вручена")
    result = ("оставлена без ответа" if c["claim_result"] == "1"
              else "ответчик отказался добровольно возместить задолженность")

    body = [
        f"1. С {d['marriage_date']} по {d['divorce_date']} я и ответчик состояли в "
        f"браке (свидетельство о заключении брака {d['marriage_cert']}). Брак "
        f"расторгнут {d['divorce_date']} ({d['divorce_cert']}).",

        f"2. В период брака, {d['contract_date']}, между мной, ответчиком и "
        f"{d['bank_name']} заключён кредитный договор № {d['contract_num']}, по "
        f"которому мы являемся созаёмщиками с солидарной ответственностью. "
        f"Обеспечение — залог квартиры по адресу: {d['apartment_address']} "
        f"(совместно нажитое имущество, ст. 34 СК РФ).",

        f"3. После расторжения брака обязательства перед банком я исполняю "
        f"единолично: все платежи вношу с моего личного счёта "
        f"№ {d['sender_account']} в {d['bank_name']}. Доля ответчика (1/2) мне не "
        f"возмещается с {d['unpaid_from']}.",

        f"4. {c['claim_sent']} в адрес ответчика направлена досудебная претензия "
        f"о возмещении 1/2 доли платежей (заказное письмо, РПО "
        f"{c['claim_track']}) {delivered}; претензия {result}. Досудебный порядок "
        f"соблюдён.",

        "5. В соответствии с п. 2 ст. 325 ГК РФ должник, исполнивший солидарную "
        "обязанность, вправе предъявить регрессное требование к остальным "
        "должникам в равных долях за вычетом доли, падающей на него самого. "
        "Общие обязательства супругов при разделе имущества распределяются "
        "пропорционально долям — по общему правилу равным (п. 1, 3 ст. 39 СК РФ). "
        "На сумму удерживаемых ответчиком денежных средств подлежат начислению "
        "проценты по ст. 395 ГК РФ (п. 48 Постановления Пленума ВС РФ от "
        "24.03.2016 № 7) по ключевой ставке Банка России, действовавшей в "
        "соответствующие периоды.",

        "6. Расчёт задолженности и процентов (по дату подачи иска):",
    ]
    for text in body:
        para(text, space_after=6)

    # Таблица расчёта
    table = doc.add_table(rows=2 + len(d["payments"]), cols=4)
    table.style = "Table Grid"
    headers = ["Дата платежа", "Сумма платежа, ₽", "1/2 доли, ₽", "Проценты ст. 395, ₽"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for i, (pdate, amount) in enumerate(d["payments"], start=1):
        table.rows[i].cells[0].text = pdate
        table.rows[i].cells[1].text = money(amount)
        table.rows[i].cells[2].text = money(amount / 2)
        table.rows[i].cells[3].text = (
            money(interest_items[i - 1][3]) if interest_items else "—")
    last = table.rows[-1].cells
    last[0].text = "ИТОГО"
    last[1].text = money(sum(a for _, a in d["payments"]))
    last[2].text = money(debt)
    last[3].text = money(interest_total) if interest_items else "расчёт прилагается"

    para("", space_after=6)

    # Просительная часть
    para("На основании изложенного, руководствуясь ст. 325, 395 ГК РФ, "
         "ст. 39 СК РФ, ст. 131–132 ГПК РФ, ПРОШУ СУД:", space_after=6)
    reqs = [
        f"1) взыскать с ответчика в мою пользу в порядке регресса "
        f"{money(debt)} ₽ — 1/2 долю внесённых мной платежей по "
        f"кредитному договору № {d['contract_num']};",
    ]
    if interest_items:
        reqs.append(
            f"2) взыскать проценты за пользование чужими денежными средствами "
            f"по ст. 395 ГК РФ на дату подачи иска в размере {money(interest_total)} ₽ "
            f"с последующим начислением на сумму основного долга по день "
            f"фактического исполнения исходя из ключевой ставки Банка России, "
            f"действующей в соответствующие периоды;")
    else:
        reqs.append(
            "2) взыскать проценты за пользование чужими денежными средствами "
            "по ст. 395 ГК РФ по день фактического исполнения исходя из ключевой "
            "ставки Банка России, действующей в соответствующие периоды "
            "(расчёт будет представлен в судебном заседании);")
    reqs.append(
        f"3) взыскать с ответчика расходы по уплате государственной пошлины "
        f"в размере {money(fee)} ₽.")
    for r in reqs:
        para(r, space_after=6)

    # Приложения
    para("Приложения (ст. 132 ГПК РФ):", bold=True, space_after=6)
    apps = [
        "1. Уведомление о вручении (направлении) копии искового заявления "
        "ответчику.",
        f"2. Квитанция об уплате государственной пошлины ({money(fee)} ₽).",
        "3. Копия кредитного договора № {} от {} с графиком платежей.".format(
            d["contract_num"], d["contract_date"]),
        f"4. Выписка из ЕГРН на квартиру ({d['apartment_address']}).",
        f"5. Копия свидетельства о заключении брака ({d['marriage_cert']}).",
        f"6. Копия свидетельства о расторжении брака ({d['divorce_cert']}).",
        "7. Выписки по счёту/платёжные документы о внесении платежей.",
        "8. Расчёт суммы иска и процентов (таблица п. 6).",
        f"9. Копия досудебной претензии с почтовой квитанцией и описью "
        f"(РПО {c['claim_track']}).",
    ]
    for a in apps:
        para(a, space_after=6)

    para("", space_after=6)
    para(f"Дата: {doc_date}        Подпись: __________ / {d['sender_name']}",
         space_after=0)

    doc.save(path)
    return price, fee


# ---------------------------------------------------------------------------
# main
# ---------------------------------------------------------------------------

def main() -> None:
    env = load_env()
    filled = sum(1 for v in env.values() if v)
    if filled:
        print(f"Загружен .env: {filled} заполненн(ых) пол(я/ей) — подаются как "
              f"умолчания.\n")
    else:
        print("Файл .env не найден или пуст — все данные запрашиваются "
              "интерактивно.\n")

    while True:
        data = collect_common(env)
        claim = collect_claim(env)
        doc_date = ask_date(
            "Дата подачи иска",
            env_default(env, "DOC_DATE", valid_date,
                        date.today().strftime("%d.%m.%Y")))

        print("\n--- Проценты по ст. 395 ГК РФ (ставка ЦБ) ---")
        interest_items, interest_total = ask_interest(
            data["payments"], doc_date, claim["claim_delivered"])

        debt = sum(a for _, a in data["payments"]) / 2
        price = debt + (interest_total or 0)
        fee = court_fee(price)
        court_type = ("мировой судья" if price <= 50_000 else "районный суд")

        court_lc = claim["court"].lower()
        if price > 50_000 and "миров" in court_lc:
            print(f"\n  ⚠ Цена иска {money(price)} ₽ превышает 50 000 ₽ — дело "
                  f"подсудно районному суду (п. 4 ч. 1 ст. 23 ГПК РФ), а не "
                  f"мировому судье! Проверьте наименование суда.")
        elif price <= 50_000 and "район" in court_lc:
            print(f"\n  ℹ Цена иска {money(price)} ₽ не превышает 50 000 ₽ — дело, "
                  f"как правило, подсудно мировому судье (ст. 23 ГПК РФ).")

        lines = [
            f"Суд:               {claim['court']} (по цене иска — {court_type})",
            f"Истец:             {data['sender_name']}, {claim['plaintiff_birth']} г.р.",
            f"Ответчик:          {data['recipient_name']}",
            f"Долг (1/2):        {money(debt)} ₽",
            f"Проценты ст. 395:  {money(interest_total)} ₽"
            if interest_items else "Проценты ст. 395:  без расчётной суммы",
            f"Цена иска:         {money(price)} ₽",
            f"Госпошлина:        {money(fee)} ₽",
            f"Дата иска:         {doc_date}",
        ]
        if confirm(lines):
            break
        print("\nНачинаем ввод заново.\n")

    safe = re.sub(r"[^\wа-яё-]+", "_", data["recipient_name"],
                  flags=re.UNICODE).strip("_")
    filename = f"Иск_{safe}_{doc_date.replace('.', '-')}.docx"
    _, fee = build_doc(data, claim, interest_items, interest_total or 0,
                       doc_date, filename)
    print(f"\nГотово: {filename}")
    print(f"Цена иска: {money(price)} ₽, госпошлина: {money(fee)} ₽ "
          f"({court_type})")


if __name__ == "__main__":
    main()
