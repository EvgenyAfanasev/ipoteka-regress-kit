#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Генератор досудебной претензии о возмещении 1/2 доли платежей по ипотеке.

Интерактивно спрашивает данные и создаёт docx-документ по шаблону из
docs/superpowers/specs/2026-08-19-regress-ipoteka-design.md (раздел 11).

Запуск:
    .venv/bin/python scripts/pretenziya.py            # интерактивный ввод
    .venv/bin/python scripts/pretenziya.py --init     # создать .env из .env.example

Значения из файла .env (в корне проекта) подставляются как умолчания:
Enter — принять значение из .env, ручной ввод — переопределить. Поля,
пустые в .env, запрашиваются интерактивно как обычно. Формат и комментарии
по каждому полю — в .env.example.
"""

import re
import sys
from datetime import date
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, Cm
except ImportError:
    sys.exit("Не найдена библиотека python-docx. Установите: pip install python-docx")


# ---------------------------------------------------------------------------
# Ввод с валидацией
# ---------------------------------------------------------------------------

def ask(prompt: str, default: str = "") -> str:
    """Обычный вопрос с необязательным значением по умолчанию."""
    suffix = f" [{default}]" if default else ""
    while True:
        raw = input(f"{prompt}{suffix}: ").strip()
        if raw:
            return raw
        if default:
            return default
        print("  Поле не может быть пустым. Попробуйте ещё раз.")


def valid_date(s: str) -> bool:
    """Формат ДД.ММ.ГГГГ и дата существует."""
    try:
        day, month, year = map(int, s.split("."))
        date(year, month, day)
        return True
    except ValueError:
        return False


def valid_money(s: str) -> bool:
    try:
        return float(s.replace(",", ".").replace(" ", "")) > 0
    except ValueError:
        return False


def valid_positive_int(s: str) -> bool:
    return s.isdigit() and int(s) > 0


def ask_date(prompt: str, default: str = "") -> str:
    """Дата в формате ДД.ММ.ГГГГ с проверкой корректности."""
    while True:
        raw = ask(prompt, default)
        if valid_date(raw):
            return raw
        print("  Неверный формат. Введите дату как ДД.ММ.ГГГГ, например 05.03.2025.")


def ask_money(prompt: str) -> float:
    """Сумма в рублях: число (запятая или точка), строго > 0."""
    while True:
        raw = ask(prompt).replace(",", ".").replace(" ", "")
        try:
            value = float(raw)
            if value <= 0:
                raise ValueError
            return value
        except ValueError:
            print("  Введите положительное число, например 38500 или 38500.50")


def ask_int(prompt: str, default: str = "") -> str:
    """Целое число > 0 (например, срок ответа в днях)."""
    while True:
        raw = ask(prompt, default)
        if valid_positive_int(raw):
            return raw
        print("  Введите целое число больше нуля, например 30")


def ask_payments(preset: list | None = None, not_before: str = "") -> list:
    """Цикл ввода платежей (с предзагрузкой из .env). Enter на пустой дате — завершение.

    not_before — дата (ДД.ММ.ГГГГ), раньше которой платёж не подлежит регрессу
    (дата расторжения брака: платежи в период брака — из общих средств, ст. 34 СК РФ).
    Платёж с более ранней датой вызывает предупреждение, но не блокируется.
    """
    def as_date(s: str):
        d, m, y = map(int, s.split("."))
        return date(y, m, d)

    payments = list(preset) if preset else []
    if payments:
        total = sum(a for _, a in payments)
        print(f"\nИз .env загружено платежей: {len(payments)}, сумма "
              f"{total:,.2f} ₽, доля 1/2 = {total / 2:,.2f} ₽.".replace(",", " "))
        print("Добавьте новые платежи или нажмите Enter для завершения.\n")
    else:
        print("\nВведите платежи по кредиту после развода, которые вы внесли единолично.")
        print("Каждый платёж: дата (ДД.ММ.ГГГГ) и сумма. Пустая дата — закончить ввод.\n")
    while True:
        raw = input("Дата платежа (Enter — закончить): ").strip()
        if not raw:
            if not payments:
                print("  Нужно хотя бы один платёж — иначе расчёт пуст.")
                continue
            break
        try:
            day, month, year = map(int, raw.split("."))
            date(year, month, day)
        except ValueError:
            print("  Неверный формат. Введите дату как ДД.ММ.ГГГГ.")
            continue
        if not_before and valid_date(not_before) and as_date(raw) < as_date(not_before):
            print(f"  ⚠ Платёж датирован раньше расторжения брака ({not_before}) — "
                  f"за период брака регресс не начисляется (ст. 34 СК РФ). "
                  f"Проверьте дату.")
        if any(p[0] == raw for p in payments):
            print("  ⚠ Платёж с этой датой уже введён — убедитесь, что это не задвоение.")
        amount = ask_money("Сумма платежа, ₽")
        payments.append((raw, amount))
        total = sum(a for _, a in payments)
        print(f"  Принято: {len(payments)} платеж(ей), сумма {total:,.2f} ₽, "
              f"доля 1/2 = {total / 2:,.2f} ₽\n".replace(",", " "))
    return payments


def confirm(summary_lines: list) -> bool:
    """Сводка всех данных + подтверждение."""
    print("\n" + "=" * 60)
    print("ПРОВЕРЬТЕ ВВЕДЁННЫЕ ДАННЫЕ".center(60))
    print("=" * 60)
    for line in summary_lines:
        print(line)
    print("=" * 60)
    while True:
        answer = input("Всё верно? (да — создать документ / нет — начать заново): ").strip().lower()
        if answer in ("да", "д", "y", "yes"):
            return True
        if answer in ("нет", "н", "n", "no"):
            return False
        print("  Ответьте «да» или «нет».")


# ---------------------------------------------------------------------------
# Файл .env
# ---------------------------------------------------------------------------

ENV_PATH = Path(__file__).resolve().parent.parent / ".env"
ENV_EXAMPLE_PATH = ENV_PATH.with_name(".env.example")


def load_env() -> dict:
    """Простой парсер KEY=VALUE без внешних зависимостей."""
    if not ENV_PATH.exists():
        return {}
    env = {}
    for line in ENV_PATH.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, _, value = line.partition("=")
        env[key.strip()] = value.strip().strip('"').strip("'")
    return env


def init_env_template() -> None:
    """Создаёт .env копированием .env.example (канонический шаблон проекта)."""
    if ENV_PATH.exists():
        sys.exit(f"{ENV_PATH} уже существует — шаблон не перезаписан. "
                 f"Отредактируйте файл вручную.")
    if not ENV_EXAMPLE_PATH.exists():
        sys.exit(f"Не найден {ENV_EXAMPLE_PATH}. Создайте .env вручную: "
                 f"cp .env.example .env")
    ENV_PATH.write_text(ENV_EXAMPLE_PATH.read_text(encoding="utf-8"),
                        encoding="utf-8")
    print(f"Создан {ENV_PATH} на основе {ENV_EXAMPLE_PATH.name}.")
    print("Заполните постоянные значения (ФИО, реквизиты, кредит) и перезапустите скрипт:")
    print("    .venv/bin/python scripts/pretenziya.py")


def parse_env_payments(raw: str):
    """'дата:сумма; дата:сумма' -> (валидные платежи, некорректные куски)."""
    payments, bad = [], []
    for chunk in raw.split(";"):
        chunk = chunk.strip()
        if not chunk:
            continue
        pdate, sep, amount = chunk.partition(":")
        pdate, amount = pdate.strip(), amount.strip()
        if not sep or not valid_date(pdate) or not valid_money(amount):
            bad.append(chunk)
            continue
        payments.append((pdate, float(amount.replace(",", ".").replace(" ", ""))))
    return payments, bad


def env_default(env: dict, key: str, validator=None, fallback: str = "") -> str:
    """Значение из .env как умолчание; пустое/невалидное -> fallback с предупреждением."""
    value = env.get(key, "")
    if not value:
        return fallback
    if validator and not validator(value):
        hint = f'использую «{fallback}»' if fallback else "введите значение вручную"
        print(f'  ⚠ {key} из .env ("{value}") не прошло проверку — {hint}.')
        return fallback
    return value


# ---------------------------------------------------------------------------
# Сбор данных
# ---------------------------------------------------------------------------

def collect_data(env: dict) -> dict:
    print("ГЕНЕРАТОР ДОСУДЕБНОЙ ПРЕТЕНЗИИ (регресс по ипотеке, п. 2 ст. 325 ГК РФ)")
    print("Значения в [скобках] — умолчания (из .env или стандартные): "
          "Enter — принять, ввод — переопределить.\n")

    d = {}
    print("--- Отправитель (вы) ---")
    d["sender_name"] = ask("Ваше ФИО полностью", env_default(env, "SENDER_NAME"))
    d["sender_address"] = ask("Ваш адрес регистрации (с индексом)",
                               env_default(env, "SENDER_ADDRESS"))
    d["sender_phone"] = ask("Ваш телефон", env_default(env, "SENDER_PHONE"))

    print("\n--- Получатель (бывшая супруга) ---")
    d["recipient_name"] = ask("ФИО полностью", env_default(env, "RECIPIENT_NAME"))
    d["recipient_address"] = ask("Адрес регистрации (с индексом)",
                                  env_default(env, "RECIPIENT_ADDRESS"))

    print("\n--- Брак ---")
    d["marriage_date"] = ask_date("Дата заключения брака",
                                  env_default(env, "MARRIAGE_DATE", valid_date))
    d["marriage_cert"] = ask("Серия и номер свидетельства о заключении брака",
                              env_default(env, "MARRIAGE_CERT"))
    d["divorce_date"] = ask_date("Дата расторжения брака",
                                  env_default(env, "DIVORCE_DATE", valid_date))
    d["divorce_cert"] = ask(
        "Свидетельство о расторжении брака (серия, номер) или суд + дата решения",
        env_default(env, "DIVORCE_CERT"))

    print("\n--- Кредитный договор ---")
    d["bank_name"] = ask("Наименование банка", env_default(env, "BANK_NAME"))
    d["contract_num"] = ask("Номер кредитного договора", env_default(env, "CONTRACT_NUM"))
    d["contract_date"] = ask_date("Дата кредитного договора",
                                   env_default(env, "CONTRACT_DATE", valid_date))
    d["sender_account"] = ask("Ваш счёт, с которого платите (№20 цифр)",
                               env_default(env, "SENDER_ACCOUNT"))
    d["apartment_address"] = ask("Адрес заложенной квартиры",
                                  env_default(env, "APARTMENT_ADDRESS"))
    d["unpaid_from"] = ask_date(
        "Дата, с которой она перестала возмещать свою долю (первый невозмещённый платёж)",
        env_default(env, "UNPAID_FROM", valid_date))

    print("\n--- Платежи ---")
    preset, bad = parse_env_payments(env.get("PAYMENTS", ""))
    if bad:
        print(f"  ⚠ В PAYMENTS из .env пропущены некорректные записи: {'; '.join(bad)}")
    d["payments"] = ask_payments(preset, not_before=d["divorce_date"])

    print("\n--- Реквизиты для возврата денег ---")
    d["refund_bank"] = ask("Банк получателя", env_default(env, "REFUND_BANK"))
    d["refund_account"] = ask("Счёт получателя", env_default(env, "REFUND_ACCOUNT"))

    d["doc_date"] = ask_date(
        "Дата претензии",
        env_default(env, "DOC_DATE", valid_date, date.today().strftime("%d.%m.%Y")))
    d["reply_days"] = ask_int("Срок ответа, дней",
                               env_default(env, "REPLY_DAYS", valid_positive_int, "30"))

    return d


def summary_of(d: dict) -> list:
    total = sum(a for _, a in d["payments"])
    lines = [
        f"Отправитель:        {d['sender_name']}, {d['sender_address']}, тел. {d['sender_phone']}",
        f"Получатель:         {d['recipient_name']}, {d['recipient_address']}",
        f"Брак:               {d['marriage_date']} ({d['marriage_cert']})",
        f"Развод:             {d['divorce_date']} ({d['divorce_cert']})",
        f"Кредит:             {d['bank_name']}, договор № {d['contract_num']} от {d['contract_date']}",
        f"Квартира:           {d['apartment_address']}",
        f"Не возмещает с:     {d['unpaid_from']}",
        f"Платежей:           {len(d['payments'])} на сумму {total:,.2f} ₽".replace(",", " "),
        f"Доля 1/2:           {total / 2:,.2f} ₽".replace(",", " "),
        f"Возврат на счёт:    {d['refund_bank']}, {d['refund_account']}",
        f"Дата претензии:     {d['doc_date']}, срок ответа: {d['reply_days']} дн.",
    ]
    return lines


# ---------------------------------------------------------------------------
# Документ
# ---------------------------------------------------------------------------

def money(value: float) -> str:
    """1250000.5 -> '1 250 000,50'"""
    return f"{value:,.2f}".replace(",", " ").replace(".", ",")


def build_doc(d: dict, path: str) -> None:
    doc = Document()

    # Базовый стиль: Times New Roman 14, полуторный интервал
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.5
    for section in doc.sections:
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    def para(text, align=None, bold=False, space_after=6):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        if align:
            p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        return p

    # Шапка (правый край)
    para(d["sender_name"], WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    para(d["sender_address"], WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    para(f"тел.: {d['sender_phone']}", WD_ALIGN_PARAGRAPH.RIGHT, space_after=12)
    para(d["recipient_name"], WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    para(d["recipient_address"], WD_ALIGN_PARAGRAPH.RIGHT, space_after=18)

    # Заголовок
    para("ДОСУДЕБНАЯ ПРЕТЕНЗИЯ", WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=0)
    para("о возмещении 1/2 доли платежей по кредитному договору",
         WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    para("(регрессное требование, п. 2 ст. 325 ГК РФ)",
         WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

    total = sum(a for _, a in d["payments"])
    half = total / 2

    body = [
        f"{d['doc_date']} я, {d['sender_name']}, направляю настоящую претензию "
        f"в связи с нижеследующим.",

        f"1. С {d['marriage_date']} по {d['divorce_date']} я и Вы состояли в браке "
        f"(свидетельство о заключении брака {d['marriage_cert']}). Брак расторгнут "
        f"{d['divorce_date']} ({d['divorce_cert']}).",

        f"2. В период брака, {d['contract_date']}, между мной, Вами и "
        f"{d['bank_name']} заключён кредитный договор № {d['contract_num']}, по которому "
        f"мы являемся созаёмщиками с солидарной ответственностью. Обеспечение — залог "
        f"квартиры по адресу: {d['apartment_address']} (совместно нажитое имущество, "
        f"ст. 34 СК РФ).",

        f"3. После расторжения брака обязательства перед банком по кредитному договору "
        f"я исполняю единолично: все платежи вношу с моего личного счёта "
        f"№ {d['sender_account']} в {d['bank_name']}. Ваша доля обязательства (1/2) мне "
        f"не возмещается с {d['unpaid_from']}.",

        "4. Согласно п. 2 ст. 325 ГК РФ должник, исполнивший солидарную обязанность, "
        "вправе предъявить регрессное требование к остальным должникам в равных долях "
        "за вычетом доли, падающей на него самого. Общие обязательства супругов при "
        "разделе имущества распределяются пропорционально долям — по общему правилу "
        "равным (п. 1, 3 ст. 39 СК РФ).",

        "5. Расчёт задолженности (платежи после расторжения брака, произведённые "
        "мной единолично):",
    ]
    for text in body:
        para(text, space_after=6)

    # Таблица платежей
    table = doc.add_table(rows=2 + len(d["payments"]), cols=3)
    table.style = "Table Grid"
    headers = ["Дата платежа", "Сумма платежа, ₽", "1/2 доли, ₽"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for i, (pdate, amount) in enumerate(d["payments"], start=1):
        table.rows[i].cells[0].text = pdate
        table.rows[i].cells[1].text = money(amount)
        table.rows[i].cells[2].text = money(amount / 2)
    last = table.rows[-1].cells
    last[0].text = "ИТОГО"
    last[1].text = money(total)
    last[2].text = money(half)

    para("", space_after=6)

    tail = [
        f"6. На основании изложенного требую в течение {d['reply_days']} календарных "
        f"дней с момента получения настоящей претензии:",

        f"1) возместить мне 1/2 долю внесённых платежей в размере {money(half)} ₽ "
        f"по реквизитам: {d['refund_bank']}, счёт {d['refund_account']}, "
        f"получатель {d['sender_name']};",

        "2) далее — возмещать 1/2 долю каждого ежемесячного платежа в течение 5 "
        "(пяти) календарных дней с даты его внесения, до полного исполнения "
        "обязательств по кредитному договору.",

        "7. В случае отказа или оставления претензии без ответа я обращусь в суд с "
        "иском о взыскании указанной суммы, а также процентов за пользование чужими "
        "денежными средствами (ст. 395 ГК РФ) и судебных расходов (госпошлина, "
        "расходы на представителя — ст. 98, 100 ГПК РФ).",
    ]
    for text in tail:
        para(text, space_after=6)

    # Приложения
    para("Приложения:", bold=True, space_after=6)
    apps = [
        f"1. Копия кредитного договора № {d['contract_num']} (или выписка по договору).",
        "2. Копии платёжных поручений / чеков по платежам из расчёта (или выписка "
        "по счёту за период).",
        "3. Расчёт задолженности (таблица п. 5).",
    ]
    for a in apps:
        para(a, space_after=6)

    # Подпись
    para("", space_after=6)
    para(f"Дата: {d['doc_date']}        Подпись: __________ / {d['sender_name']}", space_after=0)

    doc.save(path)


# ---------------------------------------------------------------------------
# main
# ---------------------------------------------------------------------------

def main() -> None:
    if "--init" in sys.argv[1:]:
        init_env_template()
        return

    env = load_env()
    filled = sum(1 for v in env.values() if v)
    if filled:
        print(f"Загружен {ENV_PATH.name}: {filled} заполненн(ых) пол(я/ей) — "
              f"подаются как умолчания.\n")
    else:
        print(f"Файл {ENV_PATH.name} не найден или пуст — все данные запрашиваются "
              f"интерактивно (создать шаблон: --init).\n")

    while True:
        data = collect_data(env)
        if confirm(summary_of(data)):
            break
        print("\nНачинаем ввод заново.\n")

    safe_name = re.sub(r"[^\wа-яё-]+", "_", data["recipient_name"], flags=re.UNICODE).strip("_")
    filename = f"Претензия_{safe_name}_{data['doc_date'].replace('.', '-')}.docx"
    build_doc(data, filename)
    total = sum(a for _, a in data["payments"])
    print(f"\nГотово: {filename}")
    print(f"Взыскание по претензии: {money(total / 2)} ₽ (1/2 от {money(total)} ₽ платежей)")


if __name__ == "__main__":
    main()
